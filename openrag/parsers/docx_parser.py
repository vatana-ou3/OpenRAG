from pathlib import Path
from typing import Any, Protocol
from uuid import uuid4

from openrag.core.document import Document
from openrag.parsers.base import BaseParser


class DocxParagraph(Protocol):
    text: str


class DocxDocument(Protocol):
    paragraphs: list[DocxParagraph]


class DocxParser(BaseParser):
    """Extract paragraph text from DOCX files using python-docx."""

    supported_extensions = {".docx"}

    def __init__(self, document_factory: Any | None = None) -> None:
        self._document_factory = document_factory

    def parse(self, path: str) -> Document:
        file_path = Path(path)
        document_factory = self._get_document_factory()
        docx_document = document_factory(str(file_path))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in docx_document.paragraphs
            if paragraph.text.strip()
        ]

        return Document(
            id=str(uuid4()),
            text="\n\n".join(paragraphs),
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
            },
        )

    def _get_document_factory(self) -> Any:
        if self._document_factory is None:
            try:
                from docx import Document as DocxDocumentFactory
            except ImportError as error:
                raise ImportError(
                    "DOCX parsing requires python-docx. Install it with "
                    "`pip install -e .[docx]` or `pip install -e .[documents]`."
                ) from error

            self._document_factory = DocxDocumentFactory

        return self._document_factory
